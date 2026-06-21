from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
PROPOSAL_DIR = ROOT / "proposals"
ASSET_DIR = ROOT / "docs" / "assets"
PROPOSAL_DIR.mkdir(exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = PROPOSAL_DIR / "english_academy_mvp_proposal_final_20260621.docx"
PNG_PATH = ASSET_DIR / "mvp-progress-20260621.png"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\malgunbd.ttf" if bold else r"C:\Windows\Fonts\malgun.ttf",
        r"C:\Windows\Fonts\NanumGothicBold.ttf" if bold else r"C:\Windows\Fonts\NanumGothic.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_multiline(draw, xy, text, fnt, fill, line_gap=8):
    x, y = xy
    for line in text.split("\n"):
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap


def create_progress_png():
    img = Image.new("RGB", (1280, 520), "#f7faf9")
    d = ImageDraw.Draw(img)
    title_f = font(34, True)
    small_f = font(15)
    h_f = font(24, True)
    p_f = font(17)
    badge_f = font(15, True)

    d.text((56, 46), "영어학원 AI 학습 대시보드 MVP 진행 현황", font=title_f, fill="#10201c")
    d.text((58, 88), "기준일: 2026-06-21 · Google Docs + Make + OpenAI + Airtable + Softr", font=small_f, fill="#52615d")

    cards = [
        ("완료", "#0f8f7b", "1단계 Airtable", "7개 테이블 구성\n샘플 학생/강사/반 입력\nLesson Feedback 저장 준비"),
        ("완료", "#0f8f7b", "2단계 자동화", "Google Docs 본문 추출\nOpenAI JSON 생성\nAirtable 저장 성공"),
        ("진행중", "#2563eb", "3단계 Softr 학생", "학생 대시보드 생성\n복습자료 상세 화면\n숙제 제출 폼 성공"),
        ("다음", "#64748b", "4단계 권한/강사", "학생별 로그인 필터\n강사 담당 반 화면\n제출물 확인/다운로드"),
    ]
    x = 56
    for badge, color, heading, body in cards:
        d.rounded_rectangle((x, 140, x + 250, 360), radius=18, fill="#ffffff", outline="#cfdad7", width=2)
        d.rounded_rectangle((x + 22, 162, x + 110, 190), radius=14, fill=color)
        d.text((x + 48 if len(badge) == 2 else x + 38, 167), badge, font=badge_f, fill="#ffffff")
        d.text((x + 24, 226), heading, font=h_f, fill="#10201c")
        draw_multiline(d, (x + 24, 266), body, p_f, "#33413d")
        x += 300

    d.rounded_rectangle((56, 406, 1206, 464), radius=12, fill="#e9f7f3", outline="#b8ded5", width=2)
    d.text(
        (80, 426),
        "현재 통과한 핵심 흐름: 강사 Google Docs 피드백 → Make/OpenAI 처리 → Airtable 저장 → Softr 학생 화면 표시 → 숙제 제출 저장",
        font=p_f,
        fill="#33413d",
    )
    img.save(PNG_PATH)


def set_normal_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Malgun Gothic"
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def create_docx():
    create_progress_png()

    doc = Document()
    set_normal_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("영어학원 AI 학습 대시보드 MVP 제안서")
    run.bold = True
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("노코드 기반 구축: Google Docs + Make.com + Airtable + Softr + OpenAI")

    for line in ["제안일: 2026년 6월 21일", "수신: 원장님", "유효기간: 제안일로부터 30일"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "1. 제안 개요")
    doc.add_paragraph(
        "안녕하세요, 원장님. 요청하신 영어학원 학습 관리 자동화 시스템을 검토하여 제안드립니다."
    )
    doc.add_paragraph(
        "이 시스템은 강사님이 수업 후 Google 문서에 작성하는 피드백을 바탕으로, AI가 학생별 복습자료, 퀴즈, 듣기 숙제, 맞춤 과제를 자동 생성하고 학생 페이지에 표시하는 구조입니다. 학생은 본인 대시보드에서 자료를 확인하고 숙제를 제출하며, 강사는 담당 반 학생의 제출물을 확인하고 다운로드할 수 있습니다."
    )
    doc.add_paragraph(
        "새 프로그램을 처음부터 개발하는 방식이 아니라, 이미 검증된 노코드 도구를 연결하는 방식이므로 빠르게 MVP를 구축하고 실제 운영 가능성을 검증할 수 있습니다."
    )

    add_heading(doc, "2. 도입 후 달라지는 점")
    add_table(
        doc,
        ["현재 수작업", "도입 후 자동화"],
        [
            ["강사가 피드백을 따로 정리해 전달", "Google 문서에 피드백만 작성하면 자동 처리"],
            ["복습자료와 숙제를 사람이 직접 제작", "AI가 학생별 복습자료와 숙제를 자동 생성"],
            ["자료를 카톡/메일로 개별 전달", "학생이 본인 대시보드에서 직접 확인"],
            ["제출 숙제를 수동 취합", "제출물이 Airtable에 자동 정리"],
            ["우수 학생 추적이 번거로움", "제출 및 활동 기록이 누적되어 관리 가능"],
        ],
    )

    add_heading(doc, "3. 주요 기능")
    for role, items in [
        ("강사", ["Google Docs 템플릿에 수업 피드백 작성", "담당 반 학생의 숙제 제출 현황 확인", "제출 파일 다운로드 및 코멘트 입력"]),
        ("학생", ["로그인 후 본인 자료만 확인", "AI 복습 요약, 단어/표현 퀴즈, 추천 듣기 영상, 맞춤 과제 확인", "숙제 답안 또는 파일 제출"]),
        ("원장/관리자", ["전체 학생, 제출 현황, 활동 기록 확인", "월간 숙제 제출 현황과 우수 학생 추적용 데이터 확보"]),
    ]:
        add_heading(doc, role, 2)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "4. 작동 방식")
    doc.add_paragraph("Google Docs 피드백 → Make.com 자동 감지 → OpenAI 분석/생성 → Airtable 저장 → Softr 학생/강사 대시보드 표시 → 학생 숙제 제출 → 강사 확인")

    add_heading(doc, "5. 현재까지 진행된 상황")
    doc.add_paragraph("현재 프로토타입 기준으로 아래 흐름까지 실제 연결 및 테스트가 완료되었습니다.")
    doc.add_picture(str(PNG_PATH), width=Inches(6.7))
    add_table(
        doc,
        ["단계", "진행 상태", "내용"],
        [
            ["1단계 Airtable", "완료", "Students, Teachers, Classes, Lesson Feedback, Homework Submissions, Listening Pool, Activity Logs 구성"],
            ["2단계 Make/OpenAI", "완료", "Google Docs 본문 추출, OpenAI JSON 생성, Airtable Lesson Feedback 저장 성공"],
            ["3단계 Softr 학생 화면", "진행 중/핵심 성공", "학생 대시보드, 복습자료 목록/상세, 숙제 제출 폼, Airtable 제출 저장 성공"],
            ["4단계 권한/강사 화면", "다음 작업", "학생별 로그인 필터, 강사 담당 반 대시보드, 제출물 다운로드/코멘트 기능 구축 예정"],
        ],
    )

    add_heading(doc, "6. 구축 범위")
    for item in [
        "Airtable 중앙 DB 설계 및 샘플 데이터 세팅",
        "Make.com 자동화: Google Docs → OpenAI → Airtable 저장",
        "Softr 학생 대시보드: 복습자료 조회 및 숙제 제출",
        "Softr 강사 대시보드: 담당 반 제출물 확인 및 다운로드",
        "기본 활동 로그 및 운영 매뉴얼",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "7. 일정 및 견적")
    doc.add_paragraph("이번 건은 첫 파일럿 구축 사례임을 감안하여, 일반 운영 포털 구축비보다 낮은 금액으로 제안드립니다.")
    add_table(
        doc,
        ["구분", "기간", "금액"],
        [["빠른 MVP", "3~4주", "200만 ~ 250만원"], ["안정화 포함", "4~5주", "280만 ~ 330만원"]],
    )
    doc.add_paragraph(
        "권장 범위는 안정화 포함 범위입니다. 학생/강사 권한 테스트, AI 결과 안정화, 예외 처리, 원장님 교육까지 포함하기 때문입니다."
    )
    doc.add_paragraph(
        "결제는 착수금 40%, 학생/강사 대시보드 완성 시 30%, 최종 인수 시 30%의 3회 분할을 권장합니다. 제안 범위 내 수정은 단계별 2회까지 무상 반영하고, 기능 추가 또는 범위 변경은 별도 협의합니다."
    )

    add_heading(doc, "8. 월 운영비 및 계정 비용")
    doc.add_paragraph(
        "아래 도구들은 원장님 또는 학원 명의 계정으로 구독하는 것을 권장합니다. 실제 비용은 계약 시점의 각 서비스 요금제와 학생 수, 사용량에 따라 달라질 수 있습니다."
    )
    add_table(
        doc,
        ["도구", "역할", "대략 비용"],
        [
            ["Make.com", "자동화 연결", "월 $9 이상 가능"],
            ["Airtable", "데이터 보관", "무료 ~ 월 $20 이상"],
            ["Softr", "학생/강사 포털", "무료 체험 후 권한/사용자 수에 따라 유료 전환 가능"],
            ["OpenAI API", "AI 생성", "사용량 과금"],
        ],
    )
    doc.add_paragraph(
        "실제 운영 단계에서는 학생/강사/관리자 권한 분리와 사용자 수에 따라 Softr 유료 플랜이 필요할 수 있습니다. 구독료와 OpenAI API 사용료는 구축비와 별도로 원장님 명의 계정에서 부담하는 방식이 적합합니다."
    )

    add_heading(doc, "9. 세금 및 정산 조건")
    doc.add_paragraph(
        "본 제안 금액은 시스템 구축 용역비 기준입니다. 제작자가 직장인 개인으로 진행하는 상황을 고려하여, 부가세, 원천징수, 기타 세무 처리상 발생하는 비용은 의뢰자 부담 또는 별도 정산을 원칙으로 합니다."
    )
    doc.add_paragraph(
        "세금계산서, 원천징수, 기타 지급 방식이 필요한 경우에는 계약 전 원장님과 협의하여 정산 방식을 확정합니다."
    )

    add_heading(doc, "10. 원장님께서 준비해 주실 것")
    for item in [
        "Google 계정 및 피드백 문서함",
        "Make, Airtable, Softr, OpenAI 계정",
        "샘플 학생/강사/반 정보",
        "추천 듣기 영상 링크 20~50개",
        "실제 운영 전 테스트 학생/강사 계정",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "11. 유지보수")
    doc.add_paragraph(
        "최종 인수 후 2주간 안정화 기간을 두어 실제 운영에서 생기는 문제를 함께 점검합니다. 이후에는 원장님 또는 담당자가 직접 운영할 수 있도록 운영 매뉴얼을 제공합니다."
    )
    doc.add_paragraph(
        "월 정액 유지보수 없이 마무리할 수 있으며, 필요 시 건별 대응 또는 월 10만~30만원 선택형 유지보수도 가능합니다."
    )

    add_heading(doc, "12. 맺음말")
    doc.add_paragraph(
        "이 시스템은 강사님의 반복 업무를 줄이고, 학생에게는 개인 맞춤 복습 환경을, 원장님께는 한눈에 보이는 관리 환경을 제공하는 것을 목표로 합니다. 작은 범위의 MVP로 빠르게 시작해 효과를 확인한 뒤 단계적으로 확장하는 방식을 권장드립니다."
    )
    doc.add_paragraph("문의나 조정이 필요한 부분은 언제든 말씀해 주시면 반영하겠습니다. 감사합니다.")
    doc.add_paragraph("최프로 드림")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    create_docx()
    print(DOCX_PATH)
    print(PNG_PATH)
